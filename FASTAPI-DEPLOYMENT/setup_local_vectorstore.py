# setup_local_vectorstore.py
"""
Local Vector Store Setup Script
- Processes PDFs/DOCX from DOCS folder
- Creates ChromaDB with embeddings
- Tracks file changes for incremental updates
- Provides progress indicators and status logging
"""
import os
import sys
import time
import json
import hashlib
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional

# Extraction / parsing
try:
    import pymupdf as fitz  # PyMuPDF for PDFs (correct import)
except ImportError:
    raise ImportError("PyMuPDF not found. Install with: pip install pymupdf")
from docx import Document           # python-docx for DOCX

# Chunking + embedding
import numpy as np
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter

# ChromaDB for local vector store
try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    print("ERROR: chromadb not installed. Install with: pip install chromadb")
    sys.exit(1)

# Load .env early
from dotenv import load_dotenv
load_dotenv()

# Configure logging with UTF-8 encoding for Windows compatibility
import sys
if sys.platform == "win32":
    # Force UTF-8 encoding for Windows console
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s • %(levelname)s • %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler("vectorstore_setup.log", encoding='utf-8')
    ]
)

# ==================== CONFIG ====================
DOCS_FOLDER = Path("DOCS")  # Relative to project root
VECTORSTORE_DIR = Path("FILES/local_vectorstore")  # Where ChromaDB will be stored
MANIFEST_FILE = VECTORSTORE_DIR / "file_manifest.json"  # Track file changes
EMBEDDING_MODEL_NAME = "sentence-transformers/all-mpnet-base-v2"
CHUNK_SIZE = 1200  # Characters (as per user request)
CHUNK_OVERLAP = 120  # Characters

# ==================== GLOBAL MODELS ====================
embedding_model: Optional[SentenceTransformer] = None
EMBED_DIM: int = None

# ==================== HELPER CLASSES ====================
class ProgressTracker:
    """Track progress with timers and status updates"""
    def __init__(self, total_files: int):
        self.total_files = total_files
        self.processed_files = 0
        self.total_chunks = 0
        self.start_time = time.time()
        self.file_start_time = None
        
    def start_file(self, filename: str):
        self.file_start_time = time.time()
        self.processed_files += 1
        logging.info(f"[{self.processed_files}/{self.total_files}] Processing: {filename}")
        
    def end_file(self, chunks_count: int):
        elapsed = time.time() - self.file_start_time if self.file_start_time else 0
        self.total_chunks += chunks_count
        logging.info(f"  [OK] Generated {chunks_count} chunks in {elapsed:.2f}s")
        
    def summary(self):
        total_elapsed = time.time() - self.start_time
        logging.info("="*60)
        logging.info(f"SETUP COMPLETE")
        logging.info(f"  Files processed: {self.processed_files}/{self.total_files}")
        logging.info(f"  Total chunks: {self.total_chunks}")
        logging.info(f"  Total time: {total_elapsed:.2f}s")
        logging.info(f"  Avg time per file: {total_elapsed/max(self.processed_files, 1):.2f}s")
        logging.info("="*60)

# ==================== FILE MANIFEST (Change Tracking) ====================
def load_manifest() -> Dict[str, Dict[str, Any]]:
    """Load file manifest for change detection"""
    if MANIFEST_FILE.exists():
        try:
            with open(MANIFEST_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logging.warning(f"Could not load manifest: {e}. Starting fresh.")
    return {}

def save_manifest(manifest: Dict[str, Dict[str, Any]]):
    """Save file manifest"""
    MANIFEST_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(MANIFEST_FILE, 'w', encoding='utf-8') as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)

def get_file_hash(filepath: Path) -> str:
    """Calculate MD5 hash of file for change detection"""
    hash_md5 = hashlib.md5()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()

def get_file_info(filepath: Path) -> Dict[str, Any]:
    """Get file metadata for change detection"""
    stat = filepath.stat()
    return {
        "hash": get_file_hash(filepath),
        "mtime": stat.st_mtime,
        "size": stat.st_size,
        "chunk_count": 0  # Will be updated after processing
    }

def should_process_file(filepath: Path, manifest: Dict[str, Dict[str, Any]]) -> bool:
    """Check if file needs processing (new or changed)"""
    file_key = str(filepath)
    if file_key not in manifest:
        return True  # New file
    
    current_info = get_file_info(filepath)
    stored_info = manifest[file_key]
    
    # Check if hash changed (file content changed)
    if current_info["hash"] != stored_info.get("hash"):
        return True
    
    # Check if modification time changed significantly
    if abs(current_info["mtime"] - stored_info.get("mtime", 0)) > 1:
        return True
    
    return False  # File unchanged

# ==================== DOCUMENT EXTRACTION ====================
def extract_table_from_pdf_page(page, page_num: int) -> tuple[List[str], List[Dict[str, Any]]]:
    """
    Extract tables from a PDF page using PyMuPDF.
    Returns: (successful_tables, failed_table_rects)
    - successful_tables: List of formatted table strings
    - failed_table_rects: List of table bounding boxes that failed (to extract as text later)
    """
    tables = []
    failed_table_rects = []
    
    try:
        # Try to find tables on the page (this can fail for some PDFs)
        tabs = page.find_tables()
        for tab_idx, tab in enumerate(tabs):
            try:
                # Get table bounding box before extraction (for fallback)
                bbox = tab.bbox if hasattr(tab, 'bbox') else None
                
                # Extract table as text
                table_text = tab.extract()
                if table_text and len(table_text) > 0:
                    # Format table as readable text
                    rows_text = []
                    for row in table_text:
                        if row:  # Skip empty rows
                            # Join row cells with | separator
                            row_text = " | ".join(str(cell).strip() if cell else "" for cell in row)
                            if row_text.strip():
                                rows_text.append(row_text)
                    if rows_text:
                        table_str = "\n".join(rows_text)
                        tables.append(f"[Table {tab_idx + 1}]:\n{table_str}")
                else:
                    # Table found but empty - mark for text extraction fallback
                    if bbox:
                        failed_table_rects.append({"bbox": bbox, "reason": "empty_table", "page": page_num})
            except Exception as e:
                # Table extraction failed - mark for text extraction fallback
                logging.warning(f"  [WARNING] Table {tab_idx + 1} on page {page_num} extraction failed: {e}. Will extract as regular text.")
                bbox = tab.bbox if hasattr(tab, 'bbox') else None
                if bbox:
                    failed_table_rects.append({"bbox": bbox, "reason": str(e), "page": page_num})
                continue
    except Exception as e:
        # If table finding fails entirely, log but continue
        logging.debug(f"Table finding failed for page {page_num}: {e}")
    
    return tables, failed_table_rects

def extract_text_from_rect(page, bbox) -> str:
    """Extract text from a specific rectangular area (fallback for failed tables)"""
    try:
        # Create a rectangle from bbox
        rect = fitz.Rect(bbox)
        # Extract text from this area
        text = page.get_text("text", clip=rect)
        return text.strip()
    except Exception as e:
        logging.debug(f"Failed to extract text from rect {bbox}: {e}")
        return ""

def extract_from_pdf(path: str) -> List[Dict[str, Any]]:
    """
    Return list of element dicts: {'text','type','page','section'}
    Includes text, headings, and tables.
    If table extraction fails, falls back to regular text extraction.
    """
    stage_start = time.time()
    logging.info("  [STAGE] Opening PDF file...")
    doc = fitz.open(path)
    logging.info(f"  [STAGE] PDF opened in {time.time() - stage_start:.2f}s ({len(doc)} pages)")
    
    elements = []
    failed_table_count = 0
    successful_table_count = 0
    
    for pno, page in enumerate(doc, start=1):
        page_start = time.time()
        
        # Extract tables first (before regular text)
        tables, failed_table_rects = extract_table_from_pdf_page(page, pno)
        if tables:
            for table_text in tables:
                elements.append({"text": table_text, "type": "table", "page": pno, "section": None})
                successful_table_count += 1
        
        # Handle failed tables - extract as regular text (don't lose content!)
        if failed_table_rects:
            failed_table_count += len(failed_table_rects)
            for failed_table in failed_table_rects:
                try:
                    # Extract text from the failed table area
                    table_text = extract_text_from_rect(page, failed_table["bbox"])
                    if table_text:
                        # Mark as table_fallback so we know it was originally a table
                        elements.append({
                            "text": f"[Table Content (extracted as text)]:\n{table_text}",
                            "type": "table_fallback",
                            "page": pno,
                            "section": None
                        })
                        logging.debug(f"  [FALLBACK] Extracted text from failed table on page {pno}")
                except Exception as e:
                    logging.warning(f"  [WARNING] Could not extract fallback text from failed table on page {pno}: {e}")
        
        # Extract regular text blocks (but exclude areas covered by tables)
        # Get table bboxes to exclude from regular text extraction
        table_bboxes = []
        try:
            tabs = page.find_tables()
            for tab in tabs:
                if hasattr(tab, 'bbox'):
                    table_bboxes.append(fitz.Rect(tab.bbox))
        except Exception:
            pass  # If table finding fails, just extract all text
        
        blocks = page.get_text("dict").get("blocks", [])
        # collect font sizes on page
        sizes = []
        for b in blocks:
            for line in b.get("lines", []):
                for span in line.get("spans", []):
                    sizes.append(span.get("size", 0))
        median_font = float(np.median(sizes)) if sizes else 0.0

        current_section = None
        for b in blocks:
            # assemble block text
            spans_text = []
            max_span_size = 0.0
            for line in b.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        spans_text.append(text)
                        max_span_size = max(max_span_size, span.get("size", 0.0))
            block_text = " ".join(spans_text).strip()
            if not block_text:
                continue
            # header heuristic
            is_header = False
            if median_font and max_span_size >= median_font * 1.15 and len(block_text.split()) <= 12:
                is_header = True

            if is_header:
                current_section = block_text
                elements.append({"text": block_text, "type": "heading", "page": pno, "section": current_section})
            else:
                elements.append({"text": block_text, "type": "paragraph", "page": pno, "section": current_section})
        
        if pno % 10 == 0:  # Log progress every 10 pages
            logging.info(f"  [STAGE] Processed {pno}/{len(doc)} pages ({time.time() - page_start:.2f}s for page {pno})")
    
    doc.close()
    extract_time = time.time() - stage_start
    logging.info(f"  [STAGE] PDF extraction complete: {len(elements)} elements in {extract_time:.2f}s")
    
    # Report table extraction statistics
    if successful_table_count > 0 or failed_table_count > 0:
        logging.info(f"  [STAGE] Table extraction: {successful_table_count} successful, {failed_table_count} failed (extracted as text)")
        if failed_table_count > 0:
            logging.warning(f"  [WARNING] {failed_table_count} tables could not be extracted in structured format and were extracted as regular text")
    
    return elements

def extract_from_docx(path: str) -> List[Dict[str, Any]]:
    """
    Return list of element dicts from DOCX: includes paragraphs, headings, and tables.
    """
    stage_start = time.time()
    logging.info("  [STAGE] Opening DOCX file...")
    doc = Document(path)
    logging.info(f"  [STAGE] DOCX opened in {time.time() - stage_start:.2f}s")
    
    elements = []
    current_section = None
    
    # Extract paragraphs and headings
    para_start = time.time()
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = getattr(para, "style", None)
        style_name = getattr(style, "name", "") if style else ""
        # consider headings
        if style_name and style_name.lower().startswith("heading"):
            current_section = text
            elements.append({"text": text, "type": "heading", "page": None, "section": current_section})
        else:
            elements.append({"text": text, "type": "paragraph", "page": None, "section": current_section})
    logging.info(f"  [STAGE] Extracted {len(elements)} paragraphs/headings in {time.time() - para_start:.2f}s")
    
    # Extract tables
    table_start = time.time()
    for table_idx, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                row_text = " | ".join(row_cells)
                rows_text.append(row_text)
        
        if rows_text:
            table_str = "\n".join(rows_text)
            elements.append({
                "text": f"[Table {table_idx + 1}]:\n{table_str}",
                "type": "table",
                "page": None,
                "section": current_section
            })
    
    if doc.tables:
        logging.info(f"  [STAGE] Extracted {len(doc.tables)} tables in {time.time() - table_start:.2f}s")
    
    logging.info(f"  [STAGE] DOCX extraction complete: {len(elements)} total elements in {time.time() - stage_start:.2f}s")
    return elements

def parse_document(filepath: str) -> List[Dict[str, Any]]:
    """Parse PDF or DOCX document with timing"""
    filepath = str(filepath)
    parse_start = time.time()
    
    if filepath.lower().endswith(".pdf"):
        logging.info(f"  [STAGE] Parsing PDF: {Path(filepath).name}")
        result = extract_from_pdf(filepath)
    elif filepath.lower().endswith(".docx"):
        logging.info(f"  [STAGE] Parsing DOCX: {Path(filepath).name}")
        result = extract_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {filepath}. Only PDF and DOCX are supported.")
    
    parse_time = time.time() - parse_start
    logging.info(f"  [STAGE] Document parsing complete in {parse_time:.2f}s")
    return result

# ==================== CHUNKING ====================
def make_chunks(elements: List[Dict[str, Any]], doc_name: str, chunk_size: int = 1200, chunk_overlap: int = 120) -> List[Dict[str, Any]]:
    """
    Combine paragraph texts into larger chunks (1200 characters with 120 overlap).
    Section/page metadata is carried into each chunk.
    Includes tables as separate elements.
    """
    chunk_start = time.time()
    logging.info(f"  [STAGE] Starting chunking (chunk_size={chunk_size}, overlap={chunk_overlap})...")
    
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    # Collect all text into one string with markers
    collect_start = time.time()
    docs_texts = []
    metas = []
    buffer = []
    last_meta = None
    table_count = 0

    for el_idx, el in enumerate(elements):
        text = el["text"].strip()
        if not text:
            continue

        # Tables are kept as separate chunks (not merged)
        # Handle both successful tables and table_fallback (failed tables extracted as text)
        if el["type"] == "table" or el["type"] == "table_fallback":
            # Flush any pending buffer first
            if buffer:
                docs_texts.append(" ".join(buffer))
                metas.append(last_meta)
                buffer = []
            # Add table as its own chunk
            meta_copy = {
                "doc_name": doc_name,
                "page": el.get("page"),
                "section": el.get("section"),
                "type": el["type"],  # Keep original type (table or table_fallback)
                "el_index": el_idx,
                "chunk_index": 0
            }
            docs_texts.append(text)
            metas.append(meta_copy)
            table_count += 1
            last_meta = None
            continue

        # If it's a heading, flush buffer and reset
        if el["type"] == "heading":
            if buffer:
                docs_texts.append(" ".join(buffer))
                metas.append(last_meta)
                buffer = []
            last_meta = {
                "doc_name": doc_name,
                "page": el.get("page"),
                "section": text,
                "type": el.get("type"),
                "el_index": el_idx,
            }
        else:
            buffer.append(text)
            last_meta = {
                "doc_name": doc_name,
                "page": el.get("page"),
                "section": el.get("section"),
                "type": el.get("type"),
                "el_index": el_idx,
            }

    # Flush last buffer
    if buffer:
        docs_texts.append(" ".join(buffer))
        metas.append(last_meta)
    
    collect_time = time.time() - collect_start
    logging.info(f"  [STAGE] Collected {len(docs_texts)} text blocks (including {table_count} tables) in {collect_time:.2f}s")

    # Now split into chunks
    split_start = time.time()
    chunks_out = []
    for meta, big_text in zip(metas, docs_texts):
        # Don't split tables or table_fallback - keep them as single chunks
        if meta.get("type") in ["table", "table_fallback"]:
            chunks_out.append({"text": big_text, "metadata": meta})
        else:
            subchunks = splitter.split_text(big_text)
            for i, sc in enumerate(subchunks):
                meta_copy = dict(meta)
                meta_copy["chunk_index"] = i
                chunks_out.append({"text": sc, "metadata": meta_copy})
    
    split_time = time.time() - split_start
    chunk_time = time.time() - chunk_start
    logging.info(f"  [STAGE] Chunking complete: {len(chunks_out)} chunks created in {chunk_time:.2f}s (split: {split_time:.2f}s)")
    return chunks_out

def clean_metadata(meta: dict) -> dict:
    """Sanitize metadata for ChromaDB (no None values)"""
    safe_meta = {}
    for k, v in meta.items():
        if v is None:
            safe_meta[k] = "unknown"
        elif isinstance(v, (str, int, float, bool)):
            safe_meta[k] = v
        elif isinstance(v, list):
            safe_meta[k] = [str(x) for x in v if x is not None]
        else:
            safe_meta[k] = str(v)
    return safe_meta

# ==================== CHROMADB OPERATIONS ====================
def clear_chromadb_collection(client: chromadb.Client, collection_name: str = "medical_documents"):
    """Clear all chunks from ChromaDB collection"""
    try:
        # Try to get existing collection
        try:
            collection = client.get_collection(name=collection_name)
            count = collection.count()
            if count > 0:
                # Delete collection and recreate
                client.delete_collection(name=collection_name)
                logging.info(f"[RESET] Deleted existing collection with {count} chunks")
        except Exception:
            # Collection doesn't exist, nothing to clear
            logging.info(f"[RESET] No existing collection found")
    except Exception as e:
        logging.warning(f"[RESET] Could not clear collection: {e}")

def init_chromadb(clear_existing: bool = False) -> chromadb.Client:
    """Initialize ChromaDB client and collection"""
    VECTORSTORE_DIR.mkdir(parents=True, exist_ok=True)
    
    client = chromadb.PersistentClient(
        path=str(VECTORSTORE_DIR),
        settings=Settings(anonymized_telemetry=False)
    )
    
    # Clear existing collection if requested
    if clear_existing:
        clear_chromadb_collection(client, "medical_documents")
    
    # Get or create collection with HNSW parameters
    # HNSW (Hierarchical Navigable Small World) parameters in metadata:
    # - hnsw:M: Maximum connections per node (default: 16, range: 4-64, higher = better recall, more memory)
    # - hnsw:construction_ef: Search width during index building (default: 100, higher = better quality, slower build)
    # - hnsw:search_ef: Search width during queries (default: 10, adjustable, higher = better recall, slower queries)
    # - hnsw:space: Distance metric (default: "l2", options: "l2", "cosine", "ip")
    try:
        # Try to get existing collection first
        collection = client.get_collection(name="medical_documents")
        logging.info(f"Using existing collection 'medical_documents' (HNSW params set at creation)")
    except Exception:
        # Create new collection with optimized HNSW parameters for medical document retrieval
        collection = client.create_collection(
            name="medical_documents",
            metadata={
                "description": "Medical document chunks with embeddings",
                # HNSW parameters optimized for medical document retrieval
                "hnsw:space": "cosine",              # Cosine similarity (better for text embeddings)
                "hnsw:M": 16,                        # Max connections per node (default: 16)
                "hnsw:construction_ef": 200,         # Build-time search width (default: 100, higher = better quality)
                "hnsw:search_ef": 10,                # Query-time search width (default: 10, adjustable)
                "hnsw:batch_size": 100,              # In-memory batch size (default: 100)
                "hnsw:sync_threshold": 1000          # Disk sync threshold (default: 1000)
            }
        )
        logging.info(f"Created new collection 'medical_documents' with HNSW config:")
        logging.info(f"  - Space: cosine")
        logging.info(f"  - M: 16 (max connections per node)")
        logging.info(f"  - construction_ef: 200 (build-time search width)")
        logging.info(f"  - search_ef: 10 (query-time search width, adjustable)")
    
    collection_count = collection.count()
    logging.info(f"ChromaDB initialized at: {VECTORSTORE_DIR} (collection has {collection_count} chunks)")
    return client, collection

def delete_file_chunks(collection: chromadb.Collection, doc_name: str):
    """Delete all chunks for a specific document"""
    try:
        # Get all IDs for this document
        results = collection.get(
            where={"doc_name": doc_name},
            include=["metadatas"]
        )
        if results["ids"]:
            collection.delete(ids=results["ids"])
            logging.info(f"  Deleted {len(results['ids'])} old chunks for {doc_name}")
    except Exception as e:
        logging.warning(f"  Could not delete old chunks for {doc_name}: {e}")

def upsert_chunks(collection: chromadb.Collection, chunks: List[Dict[str, Any]], doc_name: str, progress: ProgressTracker):
    """Embed and upsert chunks to ChromaDB with detailed timing"""
    if not chunks:
        logging.warning(f"No chunks to upsert for {doc_name}")
        return 0
    
    total = len(chunks)
    batch_size = 32
    embedded_count = 0
    num_batches = (total + batch_size - 1) // batch_size
    
    logging.info(f"  [STAGE] Starting embedding & upsert: {total} chunks in {num_batches} batches (batch_size={batch_size})...")
    embed_start = time.time()
    
    for i in range(0, total, batch_size):
        batch_num = (i // batch_size) + 1
        batch_start = time.time()
        batch = chunks[i:i+batch_size]
        texts = [c["text"] for c in batch]
        
        # Embed batch
        encode_start = time.time()
        embeddings = embedding_model.encode(texts, convert_to_numpy=True, show_progress_bar=False)
        encode_time = time.time() - encode_start
        
        # Prepare data for ChromaDB
        prep_start = time.time()
        ids = []
        metadatas = []
        documents = []
        embeddings_list = []
        
        for j, (text, emb) in enumerate(zip(texts, embeddings)):
            meta = clean_metadata(batch[j]["metadata"])
            # Create unique ID
            uid_source = f"{meta.get('doc_name','')}_{meta.get('el_index')}_{meta.get('chunk_index')}"
            vid = hashlib.md5(uid_source.encode("utf-8")).hexdigest()
            
            ids.append(vid)
            metadatas.append({
                **meta,
                "text_full": text,
                "text_snippet": text[:400]
            })
            documents.append(text)  # ChromaDB stores text separately
            embeddings_list.append(emb.tolist())
        prep_time = time.time() - prep_start
        
        # Upsert to ChromaDB
        upsert_start = time.time()
        try:
            collection.upsert(
                ids=ids,
                embeddings=embeddings_list,
                metadatas=metadatas,
                documents=documents
            )
            embedded_count += len(ids)
            upsert_time = time.time() - upsert_start
            batch_time = time.time() - batch_start
            logging.info(f"    [BATCH {batch_num}/{num_batches}] Encoded: {encode_time:.2f}s | Prepared: {prep_time:.2f}s | Upserted: {upsert_time:.2f}s | Total: {batch_time:.2f}s | Chunks: {len(ids)}")
        except Exception as e:
            logging.error(f"    [BATCH {batch_num}/{num_batches}] Failed: {e}")
            raise
    
    embed_elapsed = time.time() - embed_start
    avg_time_per_chunk = embed_elapsed / embedded_count if embedded_count > 0 else 0
    logging.info(f"  [STAGE] [OK] Embedding & upsert complete: {embedded_count} chunks in {embed_elapsed:.2f}s (avg {avg_time_per_chunk*1000:.2f}ms/chunk)")
    return embedded_count

# ==================== MAIN PROCESSING ====================
def process_file(filepath: Path, collection: chromadb.Collection, manifest: Dict[str, Dict[str, Any]], progress: ProgressTracker) -> int:
    """Process a single file with detailed timing"""
    doc_name = filepath.name
    file_key = str(filepath)
    
    progress.start_file(doc_name)
    file_start = time.time()
    
    try:
        # Parse document
        elements = parse_document(str(filepath))
        logging.info(f"  [STAGE] Extracted {len(elements)} elements from {doc_name}")
        
        # Generate chunks
        chunks = make_chunks(elements, doc_name=doc_name, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        
        # Delete old chunks for this document (if updating)
        if file_key in manifest:
            delete_start = time.time()
            delete_file_chunks(collection, doc_name)
            logging.info(f"  [STAGE] Deleted old chunks in {time.time() - delete_start:.2f}s")
        
        # Embed and upsert
        upserted = upsert_chunks(collection, chunks, doc_name, progress)
        
        # Update manifest
        file_info = get_file_info(filepath)
        file_info["chunk_count"] = len(chunks)
        file_info["processed_at"] = datetime.now().isoformat()
        manifest[file_key] = file_info
        
        file_elapsed = time.time() - file_start
        progress.end_file(len(chunks))
        logging.info(f"  [STAGE] [OK] File processing complete: {doc_name} in {file_elapsed:.2f}s total")
        return len(chunks)
        
    except Exception as e:
        logging.error(f"  [ERROR] Error processing {doc_name}: {e}", exc_info=True)
        progress.end_file(0)
        return 0

def main():
    """Main setup function"""
    import sys
    
    # Check for --reset flag
    clear_existing = "--reset" in sys.argv or "-r" in sys.argv
    
    logging.info("="*60)
    logging.info("LOCAL VECTOR STORE SETUP")
    logging.info("="*60)
    
    if clear_existing:
        logging.warning("[RESET] Clearing existing vector store (--reset flag detected)")
    
    # Check DOCS folder
    if not DOCS_FOLDER.exists():
        logging.error(f"DOCS folder not found: {DOCS_FOLDER}")
        logging.info(f"Please create {DOCS_FOLDER.absolute()} and add PDF/DOCX files")
        sys.exit(1)
    
    # Load embedding model
    logging.info(f"Loading embedding model: {EMBEDDING_MODEL_NAME}...")
    global embedding_model, EMBED_DIM
    model_start = time.time()
    embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    EMBED_DIM = embedding_model.get_sentence_embedding_dimension()
    model_elapsed = time.time() - model_start
    logging.info(f"[OK] Embedding model loaded (dim={EMBED_DIM}) in {model_elapsed:.2f}s")
    
    # Initialize ChromaDB
    logging.info("Initializing ChromaDB...")
    client, collection = init_chromadb(clear_existing=clear_existing)
    
    # Load manifest (clear if resetting)
    if clear_existing:
        manifest = {}  # Start fresh manifest when resetting
        logging.info("[RESET] Manifest cleared (starting fresh)")
    else:
        manifest = load_manifest()
        logging.info(f"Loaded manifest with {len(manifest)} tracked files")
    
    # Find all PDF/DOCX files
    all_files = sorted([f for f in DOCS_FOLDER.iterdir() if f.suffix.lower() in (".pdf", ".docx")])
    
    if not all_files:
        logging.warning(f"No PDF/DOCX files found in {DOCS_FOLDER}")
        sys.exit(1)
    
    # Filter files that need processing
    files_to_process = [f for f in all_files if should_process_file(f, manifest)]
    files_unchanged = [f for f in all_files if not should_process_file(f, manifest)]
    
    # Count by file type
    pdf_count = sum(1 for f in all_files if f.suffix.lower() == ".pdf")
    docx_count = sum(1 for f in all_files if f.suffix.lower() == ".docx")
    pdf_to_process = sum(1 for f in files_to_process if f.suffix.lower() == ".pdf")
    docx_to_process = sum(1 for f in files_to_process if f.suffix.lower() == ".docx")
    
    logging.info(f"\nFile Summary:")
    logging.info(f"  Total files found: {len(all_files)} ({pdf_count} PDF, {docx_count} DOCX)")
    logging.info(f"  Files to process: {len(files_to_process)} ({pdf_to_process} PDF, {docx_to_process} DOCX)")
    logging.info(f"  Files unchanged (skipped): {len(files_unchanged)}")
    
    if not files_to_process:
        logging.info("All files are up to date. Nothing to process.")
        return
    
    # Process files
    progress = ProgressTracker(len(files_to_process))
    
    logging.info("\n" + "="*60)
    logging.info("STARTING FILE PROCESSING")
    logging.info("="*60 + "\n")
    
    for idx, filepath in enumerate(files_to_process, 1):
        logging.info(f"\n{'='*60}")
        logging.info(f"PROCESSING FILE {idx}/{len(files_to_process)}: {filepath.name}")
        logging.info(f"{'='*60}")
        process_file(filepath, collection, manifest, progress)
        
        # Show remaining count
        remaining = len(files_to_process) - idx
        if remaining > 0:
            logging.info(f"\n[PROGRESS] {idx}/{len(files_to_process)} files completed | {remaining} remaining\n")
    
    # Save updated manifest
    save_manifest(manifest)
    logging.info(f"Manifest saved to {MANIFEST_FILE}")
    
    # Summary
    progress.summary()
    
    # Show collection stats
    collection_count = collection.count()
    logging.info(f"ChromaDB collection now contains {collection_count} total chunks")

if __name__ == "__main__":
    main()


